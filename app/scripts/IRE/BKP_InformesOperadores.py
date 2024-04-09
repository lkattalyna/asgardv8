import os
import sys
import logging
from datetime import datetime
import argparse
# ---- Librerias de Ciencia de Datos ----
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
# ---- Librerias para Generacion o Apertura de Documentos Word ----
from docx import Document
from docx.shared import RGBColor, Inches, Pt, Cm
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
from docx.enum.dml import MSO_THEME_COLOR
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

import xlsxwriter

# Configuracion de Parametros
parser = argparse.ArgumentParser()
parser.add_argument("-p", "--path", help="Ruta de los Archivos .csv")
args = parser.parse_args()

PWD = os.path.dirname(os.path.abspath(__file__))

Today = datetime.now()
Year = str(Today.year)
Month = str(Today.month)
Day = str(Today.day)
Hour = str(Today.hour)
Minute = str(Today.minute)
Second = str(Today.second)
DateA = str(Year) + '_' + str(Month) + '_' + str(Day) + '_' + str(Hour) + str(Minute) + str(Second)

DatosArchivo = []
MultiplesDates = []
document = Document()

Value_State = 0
ValueNumberFirst = 0
ValueNumberSecond = 0
MensajeCPU = ''
MensajeMEM = ''

Book = args.path + '/Archivo_Sin_Resultados.xlsx'
WorkBook = xlsxwriter.Workbook(Book)
List_Titule = [
 'HOSTNAME',
 'ESTADO']
Sheet = WorkBook.add_worksheet('Sin Resultados')
for Columns in range(0, len(List_Titule)):
    BColor = WorkBook.add_format()
    BColor.set_bg_color('24A4F7')
    BColor.set_border(1)
    BColor.set_bold()
    BColor.set_center_across()
    BColor.set_text_wrap()
    BColor.set_align('vcenter')
    Sheet.write(0, Columns, List_Titule[Columns], BColor)
    Sheet.set_column(0, Columns, 19)

def resolver_ruta(ruta_relativa):
    if hasattr(sys, '_MEIPASS'):
        return os.path.join(sys._MEIPASS, ruta_relativa)
    return os.path.join(os.path.abspath('.'), ruta_relativa)

def readFilesCsv():
    global DatosArchivo
    if args.path:
        pathCSV = args.path
        for Files in os.listdir(pathCSV):
            if Files.split('.')[(-1)].lower() == 'csv':
                logging.info(f'Leyendo Archivo: {Files}')
                FileCSV = open(pathCSV + '/' + Files, 'r')
                for linea in FileCSV.readlines():
                    DatosArchivo.append(linea.split('\\n'))
                FileCSV.close()
    else:
        raise Exception(f'Se debe ingresar una ruta para la lectura de informacion')


def makeDirectory():
    # ---- Verificacion y Creacion de Carpetas de Charts ----
    if os.path.isdir(f'{PWD}/Charts') != True:
            os.mkdir(f'{PWD}/Charts')
    # ---- Verificacion y Creacion de Carpetas de Logs ----
    if os.path.isdir(f'{PWD}/logs') != True:
            os.mkdir(f'{PWD}/logs')

class documents():
    def MainDocument(self):
        logging.info('Generando Inicio del Documento Word')
        sections = document.sections
        sections.top_margin = Cm(3.17)
        sections.bottom_margin = Cm(2.5)
        sections.left_margin = Cm(2.7)
        sections.right_margin = Cm(2.7)
        section = document.sections[0]
        header = section.header
        paragraph = header.paragraphs[0]
        Imag_header = paragraph.add_run().add_picture((resolver_ruta('Header.png')), width=(Inches(9.0)))
        paragraph.style = document.styles['Header']
        paragraph.paragraph_format.left_indent = -Inches(1.3)
        paragraph.paragraph_format.right_indent = Inches(0.5)
        section = document.sections[0]
        footer = section.footer
        paragraph = footer.paragraphs[0]
        Imag_footer = paragraph.add_run().add_picture((resolver_ruta('Footer.png')), width=(Inches(9.0)))
        paragraph.style = document.styles['Footer']
        paragraph.paragraph_format.left_indent = -Inches(1.3)
        paragraph.paragraph_format.right_indent = Inches(0.5)
        MainTitule = document.add_paragraph()
        MainTitule1 = MainTitule.add_run('\n\n\n\n\n\n\nINFORME MENSUAL {NOMBRE CLIENTE}')
        MainTitule1.font.color.rgb = RGBColor(0, 0, 0)
        MainTitule1.font.bold = True
        MainTitule1.font.size = Pt(20)
        MainTitule.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        MainParagraph = document.add_paragraph()
        MainParagraph1 = MainParagraph.add_run('\n\n\n\nEste documento contiene secretos del negocio e información de propiedad de Claro. No está permitido ningún tipo de utilización de la información contenida aquí sin previo consentimiento escrito.')
        MainParagraph1.font.color.rgb = RGBColor(141, 141, 142)
        MainParagraph1.font.size = Pt(10)
        MainParagraph.paragraph_format.left_indent = Cm(7)
        MainParagraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        document.add_page_break()
        document.add_heading('1.\tRESUMEN EJECUTIVO', level=1)
        document.add_heading('1.1.\tServidores\n', level=1)
        ListDescription = ('A nivel de sistema operativo, durante el mes reportado no se evidencian eventos extraordinarios sobre los servidores administrados.',
                            'La capacidad de los FileSystem en los servidores reportados  presentan un  porcentaje de uso normal.',
                            'Dentro de las revisiones realizadas a los logs de los servidores no se detectan alarmas y/o alertas que pongan en riesgo la estabilidad de los sistemas operativos.')
        for Description in ListDescription:
            document.add_paragraph(Description,
                style='List Bullet')
        else:
            document.add_page_break()
            document.add_heading('2.\tSistema Operativo', level=1)

    def DataDocument(self, DataMachine, CPU, MEM, PING):
        global MensajeCPU
        global MensajeMEM
        global ValueNumberFirst
        global ValueNumberSecond
        global Value_State
        machine = str(DataMachine).upper().replace('"', '')
        ValueNumberFirst = ValueNumberFirst + 1
        MinimoCPU = 100
        MaximoCPU = 0
        SumCPU = 0
        PromedioCPU = 0
        MinimoMEM = 100
        MaximoMEM = 0
        SumMEM = 0
        PromedioMEM = 0
        MinimoPING = 100
        MaximoPING = 0
        SumPING = 0
        PromedioPING = 0
        logging.info('Generando Documento con la Informacion de la Maquina: ' + machine)
        document.add_heading(('2.' + str(ValueNumberFirst) + '.\tServidor: ' + machine), level=2)
        document.add_heading(('2.' + str(ValueNumberFirst) + '.' + str(ValueNumberSecond + 1) + '.\tEstado Host\n'), level=3)
        document.add_heading(('2.' + str(ValueNumberFirst) + '.' + str(ValueNumberSecond + 2) + '.\tEstado Volumenes\n'), level=3)
        ParagraphFS = document.add_paragraph()
        ParagraphFS.add_run('Durante el mes reportado no se evidencian consumos altos a nivel de las particiones montadas en el servidor, los consumos reportados de los FileSystem existentes en el servidor son normales.')
        ParagraphFS.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        document.add_page_break()
        if os.path.isfile(PWD + '/Charts/' + machine + '_CPU.png'):
            for MPMCPU in CPU:
                if float(MPMCPU) < MinimoCPU:
                    MinimoCPU = float(MPMCPU)
                if float(MPMCPU) > MaximoCPU:
                    MaximoCPU = float(MPMCPU)
                SumCPU = SumCPU + float(MPMCPU)
                PromedioCPU = SumCPU / len(CPU)
            else:
                document.add_heading(('2.' + str(ValueNumberFirst) + '.' + str(ValueNumberSecond + 2) + '.\tConsumo CPU\n'), level=3)
                ParagraphCPU = document.add_paragraph()
                ParagraphCPU.add_run('El consumo de CPU muestra que tuvo un promedio de rendimiento del ' + str(round(PromedioCPU, 2)) + '% y un máximo del ' + str(MaximoCPU) + '%. El valor del pico no representa riesgos operativos.')
                ParagraphCPU.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                Picture = document.add_paragraph()
                Picture.add_run().add_picture((PWD + '/Charts/' + machine + '_CPU.png'), width=(Cm(12)))
                Picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
                document.add_page_break()

        else:
            Value_State = Value_State + 1
            MensajeCPU = 'No se encuentran datos de CPU para la Maquina: ' + machine
            Sheet.write(Value_State, 0, machine)
            Sheet.write(Value_State, 1, MensajeCPU)
        if os.path.isfile(PWD + '/Charts/' + machine + '_MEM.png'):
            for MPMMEM in MEM:
                if float(MPMMEM) < MinimoMEM:
                    MinimoMEM = float(MPMMEM)
                if float(MPMMEM) > MaximoMEM:
                    MaximoMEM = float(MPMMEM)
                SumMEM = SumCPU + float(MPMMEM)
                PromedioMEM = SumMEM / len(MEM)
            else:
                document.add_heading(('2.' + str(ValueNumberFirst) + '.' + str(ValueNumberSecond + 3) + '.\tConsumo Memoria\n'), level=3)
                ParagraphMemory = document.add_paragraph()
                ParagraphMemory.add_run('El consumo de Memoria muestra que tuvo un promedio de rendimiento del ' + str(round(PromedioMEM, 2)) + '% y un máximo del ' + str(MaximoMEM) + '%. El valor del pico no representa riesgos operativos.')
                ParagraphMemory.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                Picture = document.add_paragraph()
                Picture.add_run().add_picture((PWD + '/Charts/' + machine + '_MEM.png'), width=(Cm(12)))
                Picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
                document.add_page_break()

        else:
            Value_State = Value_State + 1
            MensajeMEM = 'No se encuentran datos de Memoria para la Maquina: ' + machine
            Sheet.write(Value_State, 0, machine)
            Sheet.write(Value_State, 1, MensajeMEM)
        if os.path.isfile(PWD + '/Charts/' + machine + '_PING.png'):
            for MPMPING in PING:
                if float(MPMPING) < MinimoPING:
                    MinimoPING = float(MPMPING)
                if float(MPMPING) > MaximoCPU:
                    MaximoPING = float(MPMPING)
                SumPING = SumPING + float(MPMPING)
                PromedioPING = SumPING / len(PING)
            else:
                document.add_heading(('2.' + str(ValueNumberFirst) + '.' + str(ValueNumberSecond + 4) + '.\tDisponibilidad del Servidor\n'), level=3)
                ParagraphPING = document.add_paragraph()
                ParagraphPING.add_run('El servidor presenta una disponibilidad del ' + str(round(PromedioPING, 2)) + '% para el mes reportado.')
                ParagraphPING.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                Picture = document.add_paragraph()
                Picture.add_run().add_picture((PWD + '/Charts/' + machine + '_PING.png'), width=(Cm(12)))
                Picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
                document.add_page_break()
        else:
            Value_State = Value_State + 1
            MensajePING = 'No se encuentran datos de PING para la Maquina: ' + machine
            Sheet.write(Value_State, 0, machine)
            Sheet.write(Value_State, 1, MensajePING)

class Charts:
    def CPU(self, DataMachine, dataCPU):
        machine = str(DataMachine).upper().replace('"', '')
        logging.info('Generando Grafica de CPU para la Maquina: ' + machine)
        NumberValuesCPU = []
        for CountValuesCPU in range(0, len(dataCPU)):
            if CountValuesCPU % 50 == 0:
                NumberValuesCPU.append(CountValuesCPU)
        DataFrame = pd.DataFrame(np.array(dataCPU))
        DataFrame = DataFrame.astype(float)
        plt.plot(DataFrame)
        plt.title(machine)
        plt.ylabel('VALORES CPU')
        plt.xlabel('FECHAS')
        if len(NumberValuesCPU) > len(MultiplesDates):
            NumberValuesCPU = NumberValuesCPU[:len(MultiplesDates)]
        plt.xticks(NumberValuesCPU, MultiplesDates, size=len(MultiplesDates), rotation=270)
        plt.grid()
        plt.gcf().subplots_adjust(bottom=0.4)
        plt.savefig(f'{PWD}/Charts/{machine}_CPU.png')
        plt.cla()

    def MEM(self, DataMachine, dataMEM):
        machine = str(DataMachine).upper().replace('"', '')
        logging.info('Generando Grafica de Memoria para la Maquina: ' + machine)
        NumberValuesMEM = []
        for CountValuesMEM in range(0, len(dataMEM)):
            if CountValuesMEM % 50 == 0:
                NumberValuesMEM.append(CountValuesMEM)
        DataFrame = pd.DataFrame(np.array(dataMEM))
        DataFrame = DataFrame.astype(float)
        plt.plot(DataFrame)
        plt.title(machine)
        plt.ylabel('VALORES MEM')
        plt.xlabel('FECHAS')
        if len(NumberValuesMEM) > len(MultiplesDates):
            NumberValuesMEM = NumberValuesMEM[:len(MultiplesDates)]
        plt.xticks(NumberValuesMEM, MultiplesDates, size=len(MultiplesDates), rotation=270)
        plt.grid()
        plt.gcf().subplots_adjust(bottom=0.4)
        plt.savefig(f'{PWD}/Charts/{machine}_MEM.png')
        plt.cla()

    def Ping(self, DataMachine, PING):
        machine = str(DataMachine).upper().replace('"', '')
        logging.info('Generando Grafica de Disponibilidad para la Maquina: ' + machine)
        NumberValuesPING = []
        for CountValuesPING in range(0, len(PING)):
            if CountValuesPING % 50 == 0:
                NumberValuesPING.append(CountValuesPING)
        DataFrame = pd.DataFrame(np.array(PING))
        DataFrame = DataFrame.astype(float)
        plt.plot(DataFrame)
        plt.title(machine)
        plt.ylabel('VALORES DISP')
        plt.xlabel('FECHAS')
        if len(NumberValuesPING) > len(MultiplesDates):
            NumberValuesPING = NumberValuesPING[:len(MultiplesDates)]
        plt.xticks(NumberValuesPING, MultiplesDates, size=len(MultiplesDates), rotation=270)
        plt.grid()
        plt.gcf().subplots_adjust(bottom=0.4)
        plt.savefig(f'{PWD}/Charts/{machine}_PING.png')
        plt.cla()

def analysisOfDatas():
    Machine = []
    NDate = 0
    for DataFile in range(0, len(DatosArchivo)):
        if 'CPU;' in str(DatosArchivo[DataFile]) or '"CPU";' in str(DatosArchivo[DataFile]):
            if 'CONSUMO DE CPU'.lower() not in str(DatosArchivo[DataFile]).lower():
                Machine.append(str(DatosArchivo[DataFile]).split(';')[0].replace("['", ''))
        elif 'CPU' in str(DatosArchivo[DataFile]) or '"CPU"' in str(DatosArchivo[DataFile]):
            if 'CONSUMO DE CPU'.lower() not in str(DatosArchivo[DataFile]).lower():
                if len(str(DatosArchivo[DataFile]).split('","')[0].replace("['", '').replace('"', '').split('/')) == 1:
                    Machine.append(str(DatosArchivo[DataFile]).split('","')[0].replace("['", '').replace('"', '').split('/')[0])
                else:
                    Machine.append(str(DatosArchivo[DataFile]).split('","')[0].replace("['", '').replace('"', '').split('/')[1])
        else:
            for DataFile in range(0, len(DatosArchivo)):
                if 'SiteScope Cross Performance Table' in str(DatosArchivo[DataFile]) and 'Measurement Name' in str(DatosArchivo[(DataFile + 1)]) and NDate == 0:
                    NDate = 1
                    if '","' in str(DatosArchivo[DataFile + 1]):
                        DATES = str(DatosArchivo[(DataFile + 1)]).replace("['", '').replace('"', '').replace("']", '').replace('\\n', '').split(',')[1:-1]
                    else:
                        DATES = str(DatosArchivo[(DataFile + 1)]).replace("['", '').replace('"', '').replace("']", '').replace('\\n', '').split(';')[1:-1]

    # ---- Cada Multiplos de 50 para Generar Fechas ----
    for MultiplesDate in range(0, len(DATES)):
        if MultiplesDate % 50 == 0:
            MultiplesDates.append(str(DATES[MultiplesDate]).split()[0])
    for DataMachine in Machine:
        MachineNameCPU = f"{DataMachine}/CPU/utilization".replace('"','')
        MachineNameCPUS = f"{DataMachine}/Script/CPU_Usada".replace('"','')
        MachineNameMEMS = f"{DataMachine}/Script".replace('"','')
        MachineNameMEMM = f"{DataMachine}/Memory".replace('"','')
        # ---- Segunda Opcion ----
        MachineNameMEMMS = f"{DataMachine}/Script/Memory".replace('"','')
        MachineNamePING = f"{DataMachine}/Ping".replace('"','')
        dataCPUDocument = []
        dataMEMDocument = []
        dataPINGDocument = []
        for DataFile in range(0, len(DatosArchivo)):
            if MachineNameCPU in str(DatosArchivo[DataFile]):
                if '","' in str(DatosArchivo[DataFile]):
                    MachineCPU = str(DatosArchivo[DataFile]).split('","')[1:]
                else:
                    MachineCPU = str(DatosArchivo[DataFile]).split(';')[1:]
                CPU = []
                for DataCPU in MachineCPU:
                    ValueCPU = str(DataCPU).replace('\\n', '').replace("']", '').replace(',', '.').replace('"','')
                    if ValueCPU:
                        if ValueCPU[-1] == '.':
                            ValueCPU = ValueCPU[:-1]
                        if ValueCPU == '-' or ValueCPU == '':
                            CPU.append(0)
                            dataCPUDocument.append(0)
                        else:
                            CPU.append(ValueCPU)
                            dataCPUDocument.append(ValueCPU)
                    else:
                        CPU.append(0)
                        dataCPUDocument.append(0)
                Charts().CPU(DataMachine, CPU)
            if MachineNameCPUS in str(DatosArchivo[DataFile]):
                MachineCPU = str(DatosArchivo[DataFile]).split('","')[1:]
                CPU = []
                for DataCPU in MachineCPU:
                    ValueCPU = str(DataCPU).replace('\\n', '').replace("']", '').replace(',', '.').replace('"','')
                    if ValueCPU:
                        if ValueCPU[-1] == '.':
                            ValueCPU = ValueCPU[:-1]
                        if ValueCPU == '-' or ValueCPU == '':
                            CPU.append(0)
                            dataCPUDocument.append(0)
                        else:
                            CPU.append(ValueCPU)
                            dataCPUDocument.append(ValueCPU)
                    else:
                        CPU.append(0)
                        dataCPUDocument.append(0)
                Charts().CPU(DataMachine, CPU)

            if MachineNameMEMS in str(DatosArchivo[DataFile]):
                if MachineNameMEMM in str(DatosArchivo[DataFile]):
                    if "swap" not in str(DatosArchivo[DataFile]):
                        if '","' in str(DatosArchivo[DataFile]):
                            MachineMEM = str(DatosArchivo[DataFile]).split('","')[1:]
                        else:
                            MachineMEM = str(DatosArchivo[DataFile]).split(';')[1:]
                        MEM = []
                        for DataMEM in MachineMEM:
                            ValueMEM = str(DataMEM).replace('\\n', '').replace("']", '').replace(',', '.').replace('"','')
                            if ValueMEM:
                                if ValueMEM[-1] == '.':
                                    ValueMEM = ValueMEM[:-1]
                                if ValueMEM == '-' or ValueMEM == '':
                                    MEM.append(0)
                                    dataMEMDocument.append(0)
                                else:
                                    MEM.append(ValueMEM)
                                    dataMEMDocument.append(ValueMEM)
                            else:
                                MEM.append(0)
                                dataMEMDocument.append(0)
                        Charts().MEM(DataMachine, MEM)
                if MachineNameMEMMS.lower() in str(DatosArchivo[DataFile]).lower():
                    if '","' in str(DatosArchivo[DataFile]):
                        MachineMEM = str(DatosArchivo[DataFile]).split('","')[1:]
                    else:
                        MachineMEM = str(DatosArchivo[DataFile]).split(';')[1:]
                    MEM = []
                    for DataMEM in MachineMEM:
                        ValueMEM = str(DataMEM).replace('\\n', '').replace("']", '').replace(',', '.').replace('"','')
                        if ValueMEM:
                            if ValueMEM[-1] == '.':
                                ValueMEM = ValueMEM[:-1]
                            if ValueMEM == '-' or ValueMEM == '':
                                MEM.append(0)
                                dataMEMDocument.append(0)
                            else:
                                MEM.append(ValueMEM)
                                dataMEMDocument.append(ValueMEM)
                        else:
                            MEM.append(0)
                            dataMEMDocument.append(0)
                    Charts().MEM(DataMachine, MEM)

            if MachineNamePING.lower() in str(DatosArchivo[DataFile]).lower():
                if '","' in str(DatosArchivo[DataFile]):
                    MachinePING = str(DatosArchivo[DataFile]).split('","')[1:]
                else:
                    MachinePING = str(DatosArchivo[DataFile]).split(';')[1:]
                PING = []
                for DataPING in MachinePING:
                    ValuePING = str(DataPING).replace('\\n', '').replace("']", '').replace(',', '.').replace('"','')
                    if ValuePING:
                        if ValuePING[-1] == '.':
                            ValuePING = ValuePING[:-1]
                        if ValuePING == '-' or ValuePING == '':
                            PING.append(0)
                            dataPINGDocument.append(0)
                        else:
                            PING.append(ValuePING)
                            dataPINGDocument.append(ValuePING)
                    else:
                        PING.append(0)
                        dataPINGDocument.append(0)
                Charts().Ping(DataMachine, PING)
        documents().DataDocument(DataMachine, dataCPUDocument, dataMEMDocument, dataPINGDocument)
    print('Se Genera Documento Word en la Ruta: ' + PWD + '/Informe_' + DateA + '.docx')
    print('Se Genera Documento Excel en la Ruta: ' + PWD + '/Archivo_Sin_Resultados_' + DateA + '.xlsx')
    document.save(args.path + '/Informe.docx')
    WorkBook.close()
            

# ---- Inicio de la Ejecucion ----
if __name__ == '__main__':
    makeDirectory()
    # ---- Configuracion de Nivel de Log ----
    logging.basicConfig(
        level=logging.DEBUG, 
        filename=f'{PWD}/logs/application.log',
        format = '%(asctime)-5s %(name)-15s %(levelname)-8s %(message)s'
    )
    logging.info('Iniciando Ejecucion de Generador de Informes')
    readFilesCsv()
    documents().MainDocument()
    analysisOfDatas()